import globals
import outliers
import table_data
import utils

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

document = Document()

# Title and logo
document.add_heading('Team Analysis', 0)
document.add_picture(globals.image_path, width=Inches(1.25))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Data table
stats_table = document.add_table(rows=1, cols=4)
records = table_data.init()
for key1, val1, key2, val2 in records:  
    row_cells = stats_table.add_row().cells
    row_cells[0].paragraphs[0].add_run(key1).bold = True
    row_cells[1].text = val1
    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row_cells[2].paragraphs[0].add_run(key2).bold = True
    row_cells[3].text = val2
    row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

document.add_page_break()

#Advanced stats

advanced_table = document.add_table(rows=1, cols=2)
hdr_cells = advanced_table.rows[0].cells
hdr_cells[0].text = 'Attack'
hdr_cells[0].paragraphs[0].style = 'Heading 1'
hdr_cells[1].text = 'Defense'
hdr_cells[1].paragraphs[0].style = 'Heading 1'

row_cells = advanced_table.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_picture(utils.succ_pass_toFinalThird(), height=Inches(2.2), width=Inches(2.2))
paragraph = row_cells[1].paragraphs[0]
run = paragraph.add_run()
run.add_picture(utils.succ_gkSaves(), height=Inches(2.2), width=Inches(2.2))

row_cells = advanced_table.add_row().cells
row_cells[0].text = 'Succesful passes to the final third'
row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
row_cells[1].text = 'Succesful gk saves'
row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

row_cells = advanced_table.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_picture(utils.off_duels_won(), height=Inches(2.2), width=Inches(2.2))
paragraph = row_cells[1].paragraphs[0]
run = paragraph.add_run()
run.add_picture(utils.def_duels_won(), height=Inches(2.2), width=Inches(2.2))

row_cells = advanced_table.add_row().cells
row_cells[0].text = 'Successful offensive duels'
row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
row_cells[1].text = 'Successful defensive duels'
row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

row_cells = advanced_table.add_row().cells
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_picture(utils.successful_crosses(), height=Inches(2.2), width=Inches(2.2))
row_cells[1].paragraphs[0].text = 'Crosses'
row_cells[1].paragraphs[0].style = 'Heading 1'
row_cells[1].add_paragraph()
row_cells[1].paragraphs[1].text = "According to multiple studies crosses correlate with worse results, they tend to be unreliable as shown by the graph. Team {} has executed {} out of {} aggressions.".format(globals.team_name, globals.total_crosses, globals.total_touchinbox)


document.add_page_break()

# Outliers and notes

document.add_heading('Outliers', 1)
document.add_paragraph('An outlier is a specific value out of the dataset that is "unusual". In this example an outlier will be a percentage greater than 95% or smaller than 30%')

outliers = outliers.calculate_outliers()

for key, value in outliers.items():
    document.add_paragraph(key + ' : ' + str(value) + "%")

document.add_heading('Notes', 1)

document.save('demo.docx')
